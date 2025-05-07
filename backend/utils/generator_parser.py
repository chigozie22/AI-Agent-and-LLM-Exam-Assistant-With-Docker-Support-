import io 
import pdfplumber
import pytesseract 
# import os
# import docx2txt
from PIL import Image
from docx import Document
import re

def extract_text(filename: str, content: bytes) -> str:
    ext = filename.lower().split(".")[-1]

    if ext == "pdf":
        return extract_text_from_pdf(content)
    
    elif ext == "docx":
        return extract_text_from_docx(content)
    elif ext in ["jpg", "jpeg", "png"]:
        return extract_text_from_image(content)
    else:
        raise ValueError("Unsupported file format")
    

def extract_text_from_pdf(content: bytes) -> str:
    text = ""
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            text += page.extract_text() or ""
        return text
    
def extract_text_from_docx(filename: str, content: bytes)-> str:
    doc = Document(io.BytesIO(content))
    return "\n".join([para.text] for para in doc.paragraphs)

def extract_text_from_image(content: bytes) -> str:
    image = Image.open(io.BytesIO(content))
    return pytesseract.image_to_string(image)

#extraction questions from pst question papers

def extract_questions(filename: str, content: bytes) -> str:
    ext = filename.lower().split(".")[-1]

    if ext == "pdf":
        return extract_questions_from_pdf(content)
    
    elif ext == "docx":
        return extract_questions_from_docx(content)
    elif ext in ["jpg", "jpeg", "png"]:
        return extract_questions_from_image(content)
    else:
        raise ValueError("Unsupported file format")
    
def extract_questions_from_pdf(content: bytes) -> str:
    text = ""
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            text += page.extract_text() or ""
        return text
    
def extract_questions_from_docx(content: bytes)-> str:
    doc = Document(io.BytesIO(content))
    return "\n".join([para.text] for para in doc.paragraphs)


def extract_questions_from_image(content: bytes) -> str:
    image = Image.open(io.BytesIO(content))
    raw_text = pytesseract.image_to_string(image)

    start_match = re.search(r'Question\s*1', raw_text, re.IGNORECASE)

    if start_match:
        raw_text = raw_text[start_match.start():]
    
    #normalize  whitespace
    raw_text = re.sub(r'\s+', ' ', raw_text)

    #extract questions using regex
    #use the pattern Question number
    question_blocks = re.split(r'Question\s*\d+', raw_text, flags=re.IGNORECASE)

    questions =[]
    for i in range(1, len(question_blocks), 2):
        q_number = question_blocks[i].strip()
        q_body =question_blocks[i +1].strip()
        sub_questions = re.findall(r'([abc]\.\s.*?)(?=(?:[abc]\.\s)|$)', q_body)
        full_question = [f"{q_number} {sub.strip()}" for sub in sub_questions]
        questions.extend(full_question)


    return questions